import os
import subprocess
import warnings
import time
import wave
import threading
import queue
import numpy as np
from datetime import datetime
import pytz
import streamlit as st
import yt_dlp
import json
import re
import uuid
import shutil
import atexit
from google import genai
from docx import Document

warnings.filterwarnings("ignore")

# ==========================================
# 0. Global Cleanup (Zombie Process Prevention)
# ==========================================
def cleanup_process():
    if 'app' in st.session_state and st.session_state.app['process']:
        try:
            st.session_state.app['process'].terminate()
        except:
            pass
atexit.register(cleanup_process)

# ==========================================
# 1. API Key Manager
# ==========================================
api_keys = st.secrets["gemini_api_keys"]

STATE_FILE = "api_state.json"
key_lock = threading.Lock()
global_key_index = 0

def load_key_state():
    global global_key_index
    today_str = datetime.now().strftime("%Y-%m-%d")
    if os.path.exists(STATE_FILE):
        try:
            with open(STATE_FILE, "r") as f:
                state = json.load(f)
                if state.get("date") == today_str: 
                    global_key_index = state.get("index", 0)
                else: 
                    global_key_index = 0
        except: 
            global_key_index = 0

def save_key_state(index):
    today_str = datetime.now().strftime("%Y-%m-%d")
    try:
        with open(STATE_FILE, "w") as f: 
            json.dump({"index": index, "date": today_str}, f)
    except: 
        pass

def get_next_available_client(current_index_used_by_thread):
    global global_key_index
    with key_lock:
        if current_index_used_by_thread < global_key_index:
            return genai.Client(api_key=api_keys[global_key_index]), global_key_index
        global_key_index = (global_key_index + 1) % len(api_keys)
        save_key_state(global_key_index)
        return genai.Client(api_key=api_keys[global_key_index]), global_key_index

load_key_state()

prompt = """
You are an expert, highly accurate transcriber for the Sri Lankan Parliament. 
You are transcribing a continuous live audio stream that has been split into short chunks. 

CRITICAL RULES FOR TRANSCRIPTION:

1. STRICT SCRIPT ISOLATION & ENGLISH TERMS:
   - Sinhala speech MUST ONLY be written in Sinhala script (සිංහල).
   - Tamil speech MUST ONLY be written in Tamil script (தமிழ்).
   - English speech MUST ONLY be written in English script.
   - EXCEPTION FOR TECHNICAL TERMS: If a speaker uses English terms, names, or phrases (e.g., "Brain health", "Cyber security") while speaking Sinhala, you MUST transcribe those specific words in English. DO NOT transliterate them into Sinhala letters.
   - MID-SENTENCE SWITCHING: If a speaker switches from Tamil to Sinhala, Tamil to English or Sinhala to English (or vice versa) within the same audio chunk, YOU MUST IMMEDIATELY SWITCH THE SCRIPT.
   - RESTRICTED ALPHABETS: You are STRICTLY FORBIDDEN from using any scripts other than Sinhala, Tamil, and English. NEVER use Telugu, Malayalam, Hindi, or any other characters under any circumstances.

2. NUMBERS, DATES & SYMBOLS (DICTATION FORMATTING):
   - Convert spoken numbers and dates into digits. 
   - Convert spoken punctuation/symbols into actual typographical symbols. 
   - The spoken phrase "ඇල ඉරක්" MUST be transcribed as the forward slash symbol "/".
   - The spoken phrase "බිංදුවයි" MUST be transcribed as "0".
   - Example: "අංක ආ ඇල ඉරක් විසි පහ ඇල ඉරක් බිංදුවයි අට" MUST be written as "අංක ආ/25/08".
   - Example: "දෙදහස් විසි පහේ මැයි විසි හය" MUST be written as "2025/05/26".

3. HANDLING CHUNK BOUNDARIES (CUT-OFF WORDS):
   - Because the audio is cut into chunks, boundary words might be cut in half.
   - You MUST transcribe the EXACT partial syllables/sounds you hear.
   - DO NOT guess, DO NOT hallucinate missing letters, DO NOT try to make it a valid dictionary word. Type the raw sound fragment exactly as spoken.

4. SPEAKER FOCUS (MICROPHONE ONLY) & INTERRUPTIONS:
   - Focus ONLY on the main speaker talking directly into the microphone.
   - COMPLETELY IGNORE all background yelling, side conversations, or un-mic'd voices.
   - If the chunk contains ONLY background music, noise, or silence, return an empty string "...".

FEW-SHOT EXAMPLES:
Audio: "ගරු කථානායකතුමනි, this is a serious issue regarding our brain health."
{"transcript": "ගරු කථානායකතුමනි, this is a serious issue regarding our brain health."}

Audio: "நிம்மதியாக, பாதுகாப்பாக கொண்டு நிறுத்திக்கொள்கிறேன் நன்றி. ගරු සුරංග රත්නායක මන්ත්‍රීතුමා, ඔබතුමාට විනාඩි 7ක කාලයක් ලැබෙනවා."
{"transcript": "நிம்மதியாக, பாதுகாப்பாக கொண்டு நிறுத்திக்கொள்கிறேன் நன்றி. ගරු සුරංග රත්නායක මන්ත්‍රීතුමා, ඔබතුමාට විනාඩි 7ක කාලයක් ලැබෙනවා."}

Audio: "ලිපිගොනු අංක ඒ බී ඇල ඉරක් බිංදුවයි දෙක."
{"transcript": "ලිපිගොනු අංක AB/02."}

Audio: "[incomplete sound: 'රලා වෙමින්'].. බුද්ධියෙන් අපි කටයුතු කරන්න ඕනේ..."
{"transcript": "රලා වෙමින් බුද්ධියෙන් අපි කටයුතු කරන්න ඕනේ..."}

Audio: "[background music only]"
{"transcript": "..."}

You MUST respond ONLY in the following JSON format without any markdown blocks:
{
  "transcript": "the exact text using appropriate scripts based on the rules above"
}
"""

# ==========================================
# 2. Hourly Word Auto-Save & Filters
# ==========================================
def get_hourly_filename():
    tz = pytz.timezone("Asia/Colombo")
    now = datetime.now(tz)
    date_str = now.strftime("%Y-%m-%d")
    start_hour = now.hour
    end_hour = (start_hour + 1) % 24
    return f"Transcript_{date_str}_{start_hour:02d}_to_{end_hour:02d}.docx"

def append_to_word(text):
    filename = get_hourly_filename()
    try:
        if not os.path.exists(filename):
            doc = Document()
            doc.add_heading(f"Live Transcription: {filename}", 0)
        else:
            doc = Document(filename)
        doc.add_paragraph(text)
        doc.save(filename)
    except: pass 

def remove_infinite_loops(text, max_repeats=3):
    """Removes AI hallucinated repeated words"""
    words = text.split()
    if not words: return text
    
    cleaned = [words[0]]
    count = 1
    for word in words[1:]:
        if word == cleaned[-1]:
            count += 1
            if count <= max_repeats:
                cleaned.append(word)
        else:
            count = 1
            cleaned.append(word)
    return " ".join(cleaned)

# ==========================================
# 3. Streamlit Session State (Crash-proof logic)
# ==========================================
if 'app' not in st.session_state:
    st.session_state.app = {
        'is_running': False,
        'stop_flag': [False],
        'job_queue': queue.Queue(maxsize=100), # 🔥 Limit to prevent memory leaks
        'final_transcripts': {},
        'next_chunk_to_write': 0,
        'process': None
    }
state = st.session_state.app

# ==========================================
# 4. Background Threads
# ==========================================
def worker_thread(stop_flag_list, q, transcripts_dict):
    global global_key_index
    local_client = genai.Client(api_key=api_keys[global_key_index])
    local_key_index = global_key_index

    while not stop_flag_list[0]:
        try: 
            chunk_data = q.get(timeout=1.0)
        except queue.Empty: 
            continue

        chunk_idx, chunk_filename, lk_time = chunk_data
        print(f"DEBUG: Processing chunk {chunk_idx}. File: {chunk_filename}")

        # Wait for file to stabilize (OS write delay)
        stable_size = -1
        for _ in range(15): 
            try:
                current_size = os.path.getsize(chunk_filename)
                if current_size > 44 and current_size == stable_size:
                    break 
                stable_size = current_size
            except OSError:
                pass
            time.sleep(0.1)

        # 🔥 Create Frozen Copy for safe uploading (keep extention as .wav file)
        frozen_upload_file = chunk_filename.replace(".wav", "_frozen.wav")
        try:
            shutil.copy2(chunk_filename, frozen_upload_file)
        except Exception as e:
            q.task_done()
            continue

        success = False
        retry_count = 0
        max_retries = 3

        while not success and not stop_flag_list[0] and retry_count < max_retries:
            try:
                # 🔥 FIX: කෙලින්ම MIME Type එක 'audio/wav' කියලා API එකට කියනවා
                audio_file = local_client.files.upload(file=frozen_upload_file, config={'mime_type': 'audio/wav'})
                response = local_client.models.generate_content(
                    model='gemini-3.1-flash-lite',
                    contents=[audio_file, prompt]
                )

                result_text = response.text
                
                # 🔥 Catch NoneType errors
                if not result_text:
                    raise ValueError("Empty API Response")

                cleaned_text = result_text
                json_match = re.search(r'```json\s*(.*?)\s*```', result_text, re.DOTALL)
                if json_match: 
                    cleaned_text = json_match.group(1).strip()
                else:
                    json_match_alt = re.search(r'\{.*\}', result_text, re.DOTALL)
                    if json_match_alt: 
                        cleaned_text = json_match_alt.group(0).strip()

                transcript = ""
                try:
                    result = json.loads(cleaned_text)
                    transcript = result.get("transcript", "").strip()
                except json.JSONDecodeError:
                    transcript = cleaned_text

                # 🔥 Filter out repetitive hallucinated words
                transcript = remove_infinite_loops(transcript)

                if transcript:
                    transcripts_dict[chunk_idx] = f"[{lk_time}] : {transcript}\n\n"

                local_client.files.delete(name=audio_file.name)
                success = True

            except Exception as e:
                retry_count += 1
                error_msg = str(e)
                
                # 🔥 Key Rotation logic (Includes 403 & Permission Denied)
                if any(err in error_msg for err in ["429", "Quota", "RESOURCE_EXHAUSTED", "403", "PERMISSION_DENIED"]):
                    local_client, local_key_index = get_next_available_client(local_key_index)
                    time.sleep(2)
                else:
                    time.sleep(1)

        # 🔥 Dead Letter Queue: Save file if it fails completely
        if not success:
            failed_filename = f"FAILED_{chunk_idx:03d}_{lk_time.replace(':','')}.wav"
            try:
                shutil.copy2(chunk_filename, failed_filename)
            except:
                pass
            
            transcripts_dict[chunk_idx] = f"[{lk_time} | ⚠️ MISSING AUDIO] : Failed after 3 retries. Error: {error_msg}. Saved as {failed_filename}\n\n"

        # Cleanup both original and frozen files
        try: 
            os.remove(chunk_filename)
            os.remove(frozen_upload_file)
        except: 
            pass
            
        q.task_done()

def audio_pipe_reader(pipe, stop_flag_list, q):
    audio_buffer = b""
    chunk_idx = 0
    lk_timezone = pytz.timezone("Asia/Colombo")
    
    # 🔥 Smart Chunking variables
    silence_threshold = 1000  # Increased sensitivity
    min_chunk_len = 25.0
    max_chunk_len = 60.0      # Increased to 60s to prevent word cuts

    while not stop_flag_list[0]:
        try:
            frame = pipe.read(6400)
            if not frame: 
                # 🔥 DEBUG PRINT: FFmpeg එක මැරුණොත් අල්ලගන්න
                print("🚨 DEBUG: FFmpeg stopped sending data! (Process probably crashed)")
                break
            
            if len(frame) % 2 != 0:
                frame = frame[:-(len(frame) % 2)]
            if len(frame) == 0: continue

            audio_buffer += frame
            current_duration = len(audio_buffer) / 32000.0
            audio_data = np.frombuffer(frame, dtype=np.int16)
            rms_volume = np.sqrt(np.mean(audio_data.astype(np.float32)**2)) if len(audio_data) > 0 else 0
            is_silence = rms_volume < silence_threshold

            if (current_duration >= min_chunk_len and is_silence) or (current_duration >= max_chunk_len):
                lk_time = datetime.now(lk_timezone).strftime("%H:%M:%S")
                
                # 🔥 UUID integration to prevent Zombie Thread conflicts
                unique_id = uuid.uuid4().hex[:6]
                chunk_filename = f"live_chunk_{unique_id}_{chunk_idx:03d}.wav"

                with wave.open(chunk_filename, 'wb') as wf:
                    wf.setnchannels(1); wf.setsampwidth(2); wf.setframerate(16000); wf.writeframes(audio_buffer)

                q.put((chunk_idx, chunk_filename, lk_time))
                audio_buffer = b""
                chunk_idx += 1
                
        except Exception as e:
            print(f"Audio Reader Error: {e}")
            break

# ==========================================
# 5. Streamlit UI Build
# ==========================================
st.set_page_config(layout="wide", page_title="🎙️ Gemini Array Transcriber")
st.markdown("<h1 style='text-align: center;'>🎙️ Gemini Array Real-time Transcriber</h1>", unsafe_allow_html=True)
st.caption("Auto-scaling UI with background Word Doc hourly auto-saving. Strict ordered chunk alignment.")

url = st.text_input("Enter YouTube Live URL", placeholder="https://...")

col1, col2 = st.columns(2)
# 🔥 Button lock to prevent Multiple Thread Creation
with col1: start_btn = st.button("▶ Start Transcribing", use_container_width=True, disabled=state['is_running'])
with col2: stop_btn = st.button("⏹ Stop Application", use_container_width=True)

if stop_btn:
    state['is_running'] = False
    state['stop_flag'][0] = True
    if state['process']:
        try: state['process'].terminate()
        except: pass
    st.rerun()

if start_btn and url:
    # 🔥 FFmpeg ඇත්තටම Cloud එකේ ඉන්ස්ටෝල් වෙලාද බලන අලුත් කෑල්ල
    import shutil
    if not shutil.which("ffmpeg"):
        st.error("🚨 Fatal Error: FFmpeg is NOT installed on this cloud server! Please Reboot the app from the Dashboard.")
        st.stop()

    # Reset states for a fresh run
    state['is_running'] = True
    state['stop_flag'][0] = False
    with state['job_queue'].mutex: state['job_queue'].queue.clear()
    state['final_transcripts'].clear()
    state['next_chunk_to_write'] = 0

    # Cleanup old wavs and old failed files before starting
    for f in os.listdir('.'):
        if (f.startswith("live_chunk_") or f.startswith("FAILED_")) and f.endswith(".wav"):
            try: os.remove(f)
            except: pass
        if f.endswith(".upload_ready"):
            try: os.remove(f)
            except: pass

    # Start FFmpeg
    try:
        ydl_opts = {'format': 'ba/b', 'extractor_args': {'youtube': {'player_client': ['android']}}, 'quiet': True}
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=False)
            stream_url = info.get('url', url)
    except Exception as e: # 🔥 yt-dlp Exception fixed
        st.error(f"⚠️ yt-dlp failed to extract stream! (Probably IP Blocked by YouTube). Error: {e}")
        stream_url = url

    # 🔥 DEBUG PRINT: URL එක terminal එකට දානවා
    print(f"\nDEBUG: Extracted stream URL starts with: {stream_url[:100]}...\n")

    ffmpeg_exe = './ffmpeg.exe' if os.path.exists('./ffmpeg.exe') else 'ffmpeg'
    ffmpeg_cmd = [ffmpeg_exe, "-y", "-i", stream_url, "-vn", "-acodec", "pcm_s16le", "-ar", "16000", "-ac", "1", "-f", "s16le", "-"]
    
    # 🔥 DEVNULL අයින් කරා. දැන් FFmpeg error එක terminal එකේ වැටෙනවා
    state['process'] = subprocess.Popen(ffmpeg_cmd, stdout=subprocess.PIPE)

    # Start Worker & Reader Threads
    for i in range(3):
        threading.Thread(target=worker_thread, args=(state['stop_flag'], state['job_queue'], state['final_transcripts']), daemon=True).start()
    threading.Thread(target=audio_pipe_reader, args=(state['process'].stdout, state['stop_flag'], state['job_queue']), daemon=True).start()

    st.rerun()

# ==========================================
# 6. Live Auto-Refresh Loop
# ==========================================
if state['is_running']:
    st.success("🎙️ System Live! Listening and processing chunks in order...")
    
    # Write to Word sequentially
    while state['next_chunk_to_write'] in state['final_transcripts']:
        chunk_text = state['final_transcripts'][state['next_chunk_to_write']]
        append_to_word(chunk_text.strip())
        state['next_chunk_to_write'] += 1

    # Render display text
    current_keys = list(state['final_transcripts'].keys())
    max_chunk = max(current_keys) if current_keys else -1
    
    display_text = ""
    for i in range(max_chunk + 1):
        if i in state['final_transcripts']:
            display_text += state['final_transcripts'][i]
        else:
            display_text += f"[⏳ Chunk {i} is processing...]\n\n"

    auto_scroll = st.toggle("Auto-Scroll to Bottom", value=True)

    custom_html = f"""
    <div id="transcript-box" style="
        width: 100%;
        height: 500px;
        overflow-y: auto;
        background-color: #0E1117; 
        color: #FAFAFA;
        padding: 15px;
        border-radius: 8px;
        font-family: 'Courier New', Courier, monospace;
        white-space: pre-wrap;
        border: 1px solid #4C4C54;
        line-height: 1.6;
    ">{display_text}</div>

    <script>
        var box = document.getElementById("transcript-box");
        var isAutoScroll = {'true' if auto_scroll else 'false'}; 
        var scrollPos = sessionStorage.getItem("scrollPosition");

        if (isAutoScroll) {{
            box.scrollTop = box.scrollHeight;
        }} else {{
            if (scrollPos !== null) {{
                box.scrollTop = parseInt(scrollPos);
            }}
        }}

        box.addEventListener("scroll", function() {{
            sessionStorage.setItem("scrollPosition", box.scrollTop);
        }});
    </script>
    """

    st.html(custom_html)
    
    time.sleep(2.0)
    st.rerun()
